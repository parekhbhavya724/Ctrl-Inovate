"""
Script to generate a Word Document (.docx) for SIH26189 Dataset Documentation & Rationale
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

doc = Document()

# Set Standard Page Margins (1 inch all around)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Color Palette: Navy Blue Primary, Dark Slate Secondary, Soft Background Blue
NAVY_PRIMARY = RGBColor(16, 44, 87)       # #102C57
SLATE_SECONDARY = RGBColor(53, 89, 143)   # #35598F
DARK_TEXT = RGBColor(33, 37, 41)          # #212529
BG_SHADING_HEX = "F0F4F8"

# Helper Function: Style Heading 1
def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = NAVY_PRIMARY
    return p

# Helper Function: Style Heading 2
def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = SLATE_SECONDARY
    return p

# Helper Function: Add Paragraph
def add_p(text, bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = DARK_TEXT
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.color.rgb = DARK_TEXT
    return p

# Helper Function: Set Cell Shading
def set_cell_background(cell, fill_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

# Helper Function: Callout Box
def add_callout(text, title="INVESTIGATIVE RATIONALE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, BG_SHADING_HEX)
    
    # Set left border thick navy
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="0" w:color="102C57"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>'
    tcPr.append(parse_xml(borders_xml))
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    
    r_title = p.add_run(f"💡 {title}: ")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY_PRIMARY
    
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(10.5)
    r_text.font.color.rgb = DARK_TEXT
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ---------------------------------------------------------
# DOCUMENT HEADER & TITLE SECTION
# ---------------------------------------------------------
p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_meta = p_meta.add_run("SMART INDIA HACKATHON 2026 | PROBLEM STATEMENT SIH26189")
r_meta.font.name = 'Calibri'
r_meta.font.size = Pt(9.5)
r_meta.font.bold = True
r_meta.font.color.rgb = SLATE_SECONDARY

p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(12)
p_title.paragraph_format.space_after = Pt(4)
r_title = p_title.add_run("AI-Powered Criminal Network Analysis System")
r_title.font.name = 'Calibri'
r_title.font.size = Pt(24)
r_title.font.bold = True
r_title.font.color.rgb = NAVY_PRIMARY

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(16)
r_sub = p_sub.add_run("Synthetic Dataset Architecture, Multi-Source Schema & Investigative Rationale")
r_sub.font.name = 'Calibri'
r_sub.font.size = Pt(14)
r_sub.font.italic = True
r_sub.font.color.rgb = SLATE_SECONDARY

# Organization Info Table
table_info = doc.add_table(rows=4, cols=2)
table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("Sponsoring Organization:", "Ministry of Home Affairs (MHA)"),
    ("Department:", "National Crime Records Bureau (NCRB) & Women Safety Division"),
    ("Theme & Category:", "Blockchain & Cybersecurity | Software"),
    ("Execution Security:", "100% Offline / Local Execution (Zero External Data Leakage)")
]
for idx, (label, val) in enumerate(info_data):
    cell_lbl = table_info.cell(idx, 0)
    cell_val = table_info.cell(idx, 1)
    
    cell_lbl.width = Inches(2.2)
    cell_val.width = Inches(4.3)
    
    p0 = cell_lbl.paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    r0 = p0.add_run(label)
    r0.font.bold = True
    r0.font.size = Pt(10)
    
    p1 = cell_val.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(val)
    r1.font.size = Pt(10)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ---------------------------------------------------------
# EXECUTIVE SUMMARY
# ---------------------------------------------------------
add_h1("1. Executive Summary & Problem Context")
add_p("Modern organized crime relies on complex, distributed networks spanning drug syndicates, financial money laundering rings, cyber phishing hubs, and extortion gangs. In real law enforcement scenarios, investigative data is heavily fragmented across separate units—Cyber Cells, Anti-Narcotics, Financial Intelligence Units (FIU), Telecom Providers, and Local Police Stations.")
add_p("Because real NCRB law enforcement data is classified, this project implements a Python synthetic data generator that produces a 100% fictional yet highly realistic multi-source Indian law enforcement dataset. The synthetic data deliberately embeds 4 criminal syndicates and 3 cross-network bridge connectors to evaluate AI capabilities in Named Entity Recognition (NER), Graph Construction, Centrality Analytics (PageRank/Betweenness), and Anomaly Detection.")

# ---------------------------------------------------------
# DATASET OVERVIEW TABLE
# ---------------------------------------------------------
add_h1("2. Multi-Source Dataset Architecture")
add_p("The dataset comprises 7 interconnected files saved locally in the ./data/ directory:")

tbl_overview = doc.add_table(rows=8, cols=3)
tbl_overview.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["File Name", "Record Count", "Primary Data Type & Scope"]
hdr_cells = tbl_overview.rows[0].cells
for idx, title in enumerate(headers):
    hdr_cells[idx].width = Inches(2.2) if idx != 1 else Inches(1.2)
    set_cell_background(hdr_cells[idx], "102C57")
    p = hdr_cells[idx].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)

files_summary = [
    ("entities_master.csv", "75 Profiles", "Master entity profiles (Names, Phones, Locations, Vehicles, Orgs)"),
    ("ground_truth_networks.json", "75 Mappings", "Ground-truth mapping of entity IDs to syndicates and criminal roles"),
    ("firs_reports.csv", "50 Reports", "Unstructured police narratives embedding names, numbers, and vehicles"),
    ("call_detail_records.csv", "450 CDRs", "Telecommunication logs with biased call frequency matching gang structures"),
    ("financial_transactions.csv", "220 Txns", "Banking, UPI & Crypto transfers with structuring/smurfing patterns"),
    ("social_media_intel.csv", "50 Posts", "OSINT social media feeds containing coded slang and check-ins"),
    ("criminal_history_db.csv", "20 Records", "Prior criminal court records under IPC, NDPS, and IT Act sections")
]

for row_idx, data_tuple in enumerate(files_summary, start=1):
    row_cells = tbl_overview.rows[row_idx].cells
    if row_idx % 2 == 0:
        for c in row_cells:
            set_cell_background(c, BG_SHADING_HEX)
            
    for col_idx, text in enumerate(data_tuple):
        p = row_cells[col_idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(9.5)
        if col_idx == 0:
            r.font.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ---------------------------------------------------------
# DETAILED DATASET RATIONALE SECTION
# ---------------------------------------------------------
add_h1("3. Detailed Dataset Breakdown & Investigative Rationale")

# 3.1 Entities Master
add_h2("3.1 Master Entities & Ground Truth (entities_master.csv & ground_truth_networks.json)")
add_p("Contains 75 fictional Indian profiles with fields: entity_id, name, age, gender, phone_number (+91 98xxx xxxxx), address_location, vehicle_number (MH-12-AB-1234), and known_organization.")
add_callout(
    "In graph database architectures, node resolution and entity disambiguation require fixed canonical attributes. "
    "This dataset anchors all graph nodes, while ground_truth_networks.json enables hackathon evaluators to programmatically measure Precision, Recall, and F1-Score for AI node clustering.",
    "INVESTIGATIVE RATIONALE"
)

# 3.2 FIR Reports
add_h2("3.2 Police FIR Narratives (firs_reports.csv)")
add_p("Contains 50 narrative police reports from stations such as Cyber Crime PS Cyberabad and Ratanpur Crime Branch, covering offences like Cyber Fraud, Narcotics Smuggling, Hawala, and Arms Possession.")
add_callout(
    "Real investigative data begins as freeform unstructured text. This dataset tests the Natural Language Processing (NLP) "
    "and Named Entity Recognition (NER) pipeline's ability to automatically extract PERSON, PHONE, VEHICLE, LOCATION, and ORGANIZATION entities without human manual tagging.",
    "INVESTIGATIVE RATIONALE"
)

# 3.3 CDRs
add_h2("3.3 Call Detail Records (call_detail_records.csv)")
add_p("Contains 450 call records specifying caller_id, callee_id, timestamp, duration_seconds, and cell_tower_location.")
add_callout(
    "CDR analysis forms the backbone of law enforcement link analysis. Intra-network members call each other with high frequency (65%), "
    "cross-network bridges connect different syndicates (15%), and background civilian calls (20%) test the AI's noise filtering capabilities.",
    "INVESTIGATIVE RATIONALE"
)

# 3.4 Financial Transactions
add_h2("3.4 Financial Transactions Ledger (financial_transactions.csv)")
add_p("Contains 220 transaction logs covering UPI, Cash Deposits, RTGS/NEFT, and Crypto Transfers with amounts and timestamps.")
add_callout(
    "Following the money trail uncovers high-level ringleaders who avoid direct phone calls or physical crimes. "
    "This dataset specifically embeds Financial Structuring ('Smurfing')—multiple transfers just under ₹50,000 to bypass FIU alerts—and multi-hop Hawala layering.",
    "INVESTIGATIVE RATIONALE"
)

# 3.5 Social Media Intel
add_h2("3.5 OSINT Social Media Intelligence (social_media_intel.csv)")
add_p("Contains 50 synthetic posts across platforms like Chatgram, DarkPost, and X-Feed featuring coded language and location mentions.")
add_callout(
    "Modern gangs utilize encrypted platforms and social channels for operational coordination. This dataset tests handle-to-entity alias matching "
    "and keyword sentiment anomaly detection (e.g., 'Package arrived at warehouse', 'Transfer settled on UPI').",
    "INVESTIGATIVE RATIONALE"
)

# 3.6 Criminal History
add_h2("3.6 Criminal History Database (criminal_history_db.csv)")
add_p("Contains 20 prior legal conviction and case records under IPC, NDPS Act, Arms Act, and IT Act sections.")
add_callout(
    "In law enforcement threat scoring matrices, nodes with prior convictions under violent or narcotic sections receive an elevated risk score "
    "when combined with active CDR call spikes and financial layering activity.",
    "INVESTIGATIVE RATIONALE"
)

# ---------------------------------------------------------
# GROUND TRUTH NETWORKS BREAKDOWN
# ---------------------------------------------------------
add_h1("4. Ground-Truth Criminal Syndicates & Graph Validation")
add_p("To benchmark PageRank (Kingpin detection), Betweenness Centrality (Bridge identification), and Louvain Community Detection (Cell clustering), the dataset embeds 4 hidden syndicates:")

tbl_nets = doc.add_table(rows=5, cols=4)
tbl_nets.alignment = WD_TABLE_ALIGNMENT.CENTER

net_headers = ["Network ID", "Syndicate Domain", "Member Count", "Kingpin / Leader"]
for idx, title in enumerate(net_headers):
    cell = tbl_nets.rows[0].cells[idx]
    set_cell_background(cell, "102C57")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)

syndicate_data = [
    ("NET_ALPHA", "Hawala & Crypto Money Laundering Ring", "8 Members", "ENT_001"),
    ("NET_BETA", "Narcotics & Contraband Smuggling Cartel", "8 Members", "ENT_009"),
    ("NET_GAMMA", "Cyber Fraud & Phishing Syndicate", "8 Members", "ENT_016"),
    ("NET_DELTA", "Illegal Firearms & Extortion Gang", "6 Members", "ENT_023")
]

for row_idx, data_tuple in enumerate(syndicate_data, start=1):
    row_cells = tbl_nets.rows[row_idx].cells
    if row_idx % 2 == 0:
        for c in row_cells:
            set_cell_background(c, BG_SHADING_HEX)
            
    for col_idx, text in enumerate(data_tuple):
        p = row_cells[col_idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(9.5)
        if col_idx in [0, 3]:
            r.font.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(10)

add_h2("4.1 Cross-Network Bridge Connectors (High Betweenness Centrality)")
add_p("ENT_007 (Hawala Conduit): Belongs to NET_ALPHA and NET_BETA. Launders drug money through Hawala channels.", bold_prefix="• ")
add_p("ENT_015 (Identity Supplier): Belongs to NET_BETA and NET_GAMMA. Supplies stolen IDs from drug victims to cyber fraud rings.", bold_prefix="• ")
add_p("ENT_022 (Extortion Financier): Belongs to NET_GAMMA and NET_DELTA. Uses cyber fraud proceeds to fund illegal arms procurement for extortion gangs.", bold_prefix="• ")

# ---------------------------------------------------------
# COMPLIANCE & VERIFICATION
# ---------------------------------------------------------
add_h1("5. Ethics, Privacy & Local Execution Compliance")
add_p("100% Synthetic Data: All names, numbers, addresses, and narrative scenarios are completely fictional generated via Faker (en_IN) with fixed random seed 42.", bold_prefix="1. ")
add_p("Local Execution Guarantee: As requested, the generation script and output data reside 100% locally on your machine. No data or code was pushed to remote Git repositories.", bold_prefix="2. ")
add_p("Reproducibility: The generator script generate_dataset.py can be re-run at any time to regenerate the exact dataset.", bold_prefix="3. ")

# Save Document
output_docx_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'SIH26189_Dataset_Documentation_and_Rationale.docx')
doc.save(output_docx_path)
print(f"[OK] Word Document successfully created at: {output_docx_path}")
